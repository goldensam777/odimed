import os
import re
import uuid
from typing import Annotated, Any

from docx import Document
from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from sqlmodel import select

from app.api.deps import CurrentUser, SessionDep
from app.models import (
    OrdonnanceTemplate,
    OrdonnanceTemplatePublic,
    ProfilMedecin,
)

router = APIRouter(prefix="/templates", tags=["templates"])

UPLOAD_DIR_TEMPLATES = "uploads/templates"


def extract_docx_tokens(file_path: str) -> list[str]:
    """
    Extrait tous les tokens sous la forme $nom_token$ ou $nom_token:img$ du document Word.
    """
    try:
        doc = Document(file_path)
        full_text = []
        for p in doc.paragraphs:
            full_text.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        combined = " ".join(full_text)
        tokens = re.findall(r"\$[a-zA-Z0-9_:]+\$", combined)
        return list(set(tokens))
    except Exception:
        return []


@router.get("/", response_model=list[OrdonnanceTemplatePublic])
def list_my_templates(session: SessionDep, current_user: CurrentUser) -> Any:
    """
    Lister tous les modèles d'ordonnances (.docx) enregistrés par le médecin connecté.
    """
    statement_medecin = select(ProfilMedecin).where(ProfilMedecin.user_id == current_user.id)
    profil = session.exec(statement_medecin).first()
    if not profil:
        return []

    statement = select(OrdonnanceTemplate).where(OrdonnanceTemplate.medecin_id == profil.id)
    templates = session.exec(statement).all()
    return templates


@router.post("/upload", response_model=dict[str, Any])
def upload_template(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    nom_template: Annotated[str, Form(...)],
    file: Annotated[UploadFile, File(...)],
) -> Any:
    """
    Uploader un fichier modèle .docx et détecter automatiquement les tokens ($patient$, $signature:img$, etc.)
    """
    statement_medecin = select(ProfilMedecin).where(ProfilMedecin.user_id == current_user.id)
    profil = session.exec(statement_medecin).first()
    if not profil:
        raise HTTPException(
            status_code=400,
            detail="Vous devez créer un profil médecin avant d'uploader un modèle d'ordonnance.",
        )

    if not file.filename or not file.filename.endswith(".docx"):
        raise HTTPException(
            status_code=400, detail="Seuls les fichiers Microsoft Word (.docx) sont acceptés."
        )

    os.makedirs(UPLOAD_DIR_TEMPLATES, exist_ok=True)
    filename = f"{profil.id}_{uuid.uuid4().hex[:8]}.docx"
    file_path = os.path.join(UPLOAD_DIR_TEMPLATES, filename)

    with open(file_path, "wb") as f:
        f.write(file.file.read())

    tokens = extract_docx_tokens(file_path)

    template_obj = OrdonnanceTemplate(
        medecin_id=profil.id,
        nom_template=nom_template,
        chemin_fichier_docx=file_path,
    )
    session.add(template_obj)
    session.commit()
    session.refresh(template_obj)

    return {
        "template": OrdonnanceTemplatePublic.model_validate(template_obj),
        "detected_tokens": tokens,
    }


@router.delete("/{id}")
def delete_template(session: SessionDep, current_user: CurrentUser, id: uuid.UUID) -> Any:
    """
    Supprimer un modèle d'ordonnance.
    """
    statement_medecin = select(ProfilMedecin).where(ProfilMedecin.user_id == current_user.id)
    profil = session.exec(statement_medecin).first()
    if not profil:
        raise HTTPException(status_code=403, detail="Permissions insuffisantes.")

    template_obj = session.get(OrdonnanceTemplate, id)
    if not template_obj or template_obj.medecin_id != profil.id:
        raise HTTPException(status_code=404, detail="Modèle introuvable.")

    if os.path.exists(template_obj.chemin_fichier_docx):
        try:
            os.remove(template_obj.chemin_fichier_docx)
        except OSError:
            pass

    session.delete(template_obj)
    session.commit()
    return {"message": "Modèle supprimé avec succès."}
