import io

from docx import Document
from fastapi.testclient import TestClient

from app.core.config import settings


def create_sample_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("Ordonnance Médicale", level=1)
    doc.add_paragraph("Patient: $nom_patient$")
    doc.add_paragraph("Prescription: $posologie$")
    doc.add_paragraph("Signature: $signature:img$")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def test_upload_and_list_templates(
    client: TestClient, normal_user_token_headers: dict[str, str]
) -> None:
    # 1. Créer le profil médecin d'abord
    client.post(
        f"{settings.API_V1_STR}/medecins/me",
        headers=normal_user_token_headers,
        json={
            "numero_ordre": "MED-TEST-TMPL",
            "specialite": "Cardiologie",
            "pays_exercice": "Bénin",
        },
    )

    # 2. Upload modèle .docx
    docx_bytes = create_sample_docx_bytes()
    files = {"file": ("template_test.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    data = {"nom_template": "Modèle Standard Cardiologie"}

    response = client.post(
        f"{settings.API_V1_STR}/templates/upload",
        headers=normal_user_token_headers,
        data=data,
        files=files,
    )
    assert response.status_code == 200
    res_json = response.json()
    assert "detected_tokens" in res_json
    assert "$nom_patient$" in res_json["detected_tokens"]
    assert "$signature:img$" in res_json["detected_tokens"]

    # 3. Lister les modèles
    list_res = client.get(
        f"{settings.API_V1_STR}/templates/",
        headers=normal_user_token_headers,
    )
    assert list_res.status_code == 200
    assert len(list_res.json()) >= 1
